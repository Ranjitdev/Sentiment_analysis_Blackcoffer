from src.exception import CustomException
from src.logger import logging
import pandas as pd
from docx import Document
from dataclasses import dataclass
import os
import sys


@dataclass
class DataIngesionConfig:
    notebook_input = 'notebook/input.xlsx'
    output_format = r'notebook/output.xlsx'
    objective_file = r'notebook/objective.docx'
    test_analysis = r'notebook/test_analysis.docx'
    output_file = r'artifacts/output.csv'
    input_file = r'artifacts/input.csv'


class DataIngesion:
    def __init__(self):
        self.config = DataIngesionConfig
        os.makedirs(os.path.dirname(self.config.output_file), exist_ok=True)
        files =  os.listdir('artifacts')
        if 'input.csv' not in files:
            in_file = pd.read_excel(self.config.notebook_input, index_col=None)
            in_file.to_csv(self.config.input_file, index=False)
            logging.info('Input file ingested')
        elif 'output.csv' not in files:
            out_file = pd.read_excel(self.config.output_format, index_col=None)
            out_file.to_csv(self.config.output_file, index=False)
            logging.info('Output file ingested')


    def get_data(self, content: str):
        try:
            if content == 'Websites List':
                return pd.read_csv(self.config.input_file, index_col=None)
            elif content == 'output format':
                return pd.read_excel(self.config.output_format, index_col=None)
            elif content == 'Website Contents':
                return pd.read_csv(self.config.output_file, index_col=None)
            elif content == 'objective':
                obj = Document(self.config.objective_file)
                text = []
                for para in obj.paragraphs:
                    text.append(para.text)
                return text
            elif content == 'test analysis':
                obj = Document(self.config.test_analysis)
                text = []
                for para in obj.paragraphs:
                    text.append(para.text)
                return ' '.join(text)
        except Exception as e:
            raise CustomException(e, sys)
